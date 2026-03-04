import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

SRC  = r"C:\Projekte\MoPilot\MoPilot_Prototype_v0.1.docx"
DEST = r"C:\Projekte\MoPilot\MoPilot_Prototype_v0.2.docx"

doc = Document(SRC)

# ─── Helpers ───────────────────────────────────────────────────────────────

def add_h1(doc, text):
    p = doc.add_paragraph(text, style='Heading 1')
    return p

def add_h2(doc, text):
    p = doc.add_paragraph(text, style='Heading 2')
    return p

def add_h3(doc, text):
    p = doc.add_paragraph(text, style='Heading 3')
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text, style='Normal')
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_mono(doc, text):
    """Code-like monospace paragraph."""
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    p.paragraph_format.left_indent = Cm(1)
    return p

def add_box(doc, lines):
    """ASCII-box style like existing doc."""
    for line in lines:
        add_mono(doc, line)

# ─── Version + Datum auf Titelseite aktualisieren ──────────────────────────
for para in doc.paragraphs:
    if 'Prototyp v0.1' in para.text:
        for run in para.runs:
            if 'v0.1' in run.text:
                run.text = run.text.replace('v0.1', 'v0.2')
    if '28. Februar 2026' in para.text:
        for run in para.runs:
            if '28. Februar 2026' in run.text:
                run.text = run.text.replace('28. Februar 2026', '1. Maerz 2026')
    # Volltext-Fallback
    if para.text.strip() == 'Prototyp v0.1':
        para.clear()
        para.add_run('Prototyp v0.2')

# ─── Inhaltsverzeichnis ergaenzen ──────────────────────────────────────────
for para in doc.paragraphs:
    if '4.  Naechste Schritte' in para.text:
        para.clear()
        run = para.add_run('4.  Naechste Schritte')
        para.add_run('\n5.  ZEO Kundenassistent – Entwicklung und Deployment (zeo-kunden.mopilot.website)')
        para.add_run('\n5.1  Projektbeschreibung und Ziel')
        para.add_run('\n5.2  Technische Architektur')
        para.add_run('\n5.3  Entwicklungsprozess mit Claude Code')
        para.add_run('\n5.4  Deployment-Ablauf: Herausforderungen und Loesungen')
        para.add_run('\n5.5  Live-System und Ergebnis')
        para.add_run('\n6.  Naechstes Sub-Projekt: hotline.mopilot.website')
        para.add_run('\n6.1  Konzept und Zielgruppe')
        para.add_run('\n6.2  Technische Anforderungen')
        para.add_run('\n6.3  Entwicklungsschritte')
        para.add_run('\n6.4  Deployment-Checkliste')
        para.add_run('\n6.5  Zeitplan')
        break

# ─── Kapitel 5: ZEO Kundenassistent ───────────────────────────────────────

add_h1(doc, '5. ZEO Kundenassistent – Entwicklung und Deployment')
add_body(doc,
    'Im Rahmen des MoPilot-Projekts wurde ein eigenstaendiger, KI-gestuetzter Kundenassistent fuer '
    'ZEO E-Carsharing entwickelt und unter https://zeo-kunden.mopilot.website in Betrieb genommen. '
    'Dieses Kapitel dokumentiert den vollstaendigen Prozess – von der Anforderungsanalyse bis zum '
    'produktiven Betrieb – und zeigt, wie die vorhandene MoPilot-Infrastruktur wiederverwendet wurde.')

# 5.1
add_h2(doc, '5.1 Projektbeschreibung und Ziel')
add_body(doc,
    'Waehrend MoPilot (mopilot.website) ein rollenbasiertes Demo-System fuer verschiedene '
    'Carsharing-Akteure ist, richtet sich der ZEO Kundenassistent ausschliesslich an Endkunden '
    'von ZEO Carsharing. Das Ziel war ein oeffentlich zugaengliches, informatives und '
    'KI-gestuetztes Tool, das ohne Login nutzbar ist und alle ZEO-relevanten Fragen beantwortet.')

add_body(doc, 'Die App besteht aus fuenf interaktiven Modulen auf einer Single-Page:')
add_bullet(doc, 'KI-Chat (Hero-Element): SSE-Streaming-Assistent mit ZEO-Systemkontext und Beispielfragen')
add_bullet(doc, 'Modul 1 – Kostenrechner: Drei Schieberegler errechnen den optimalen Tarif (eco vs. eco plus)')
add_bullet(doc, 'Modul 2 – ÖPNV-Karte: Leaflet/OpenStreetMap mit 12 ZEO-Stationen und ÖPNV-Anbindung')
add_bullet(doc, 'Modul 3 – Reichweiten-Guide: Temperaturabhaengige Reichweitenberechnung fuer alle 9 Fahrzeugmodelle')
add_bullet(doc, 'Modul 4 – Fahrzeugkatalog: Filterbare Kacheln mit Specs, Eignung und Reichweite')
add_bullet(doc, 'Modul 5 – Onboarding: 6-Schritt-Guide und FAQ-Akkordeon mit 18 haeufigen Fragen')

# 5.2
add_h2(doc, '5.2 Technische Architektur')
add_body(doc,
    'Die Architektur wurde bewusst einfach gehalten: keine Datenbank, kein Authentifizierungssystem, '
    'nur das Noetigste fuer Chat und statische Inhalte. Alle ZEO-Inhalte (Tarife, Fahrzeuge, '
    'Stationen, FAQs) liegen als gepflegte JSON-Dateien im Repository.')

add_box(doc, [
    '+============================================================+',
    '|               NUTZER (Browser)                             |',
    '|   https://zeo-kunden.mopilot.website                       |',
    '+==========================+=================================+',
    '                           |  HTTPS',
    '                           v',
    '+============================================================+',
    '|         HETZNER CX33  (142.132.232.211)                    |',
    '|                                                            |',
    '|  +--------------------------------------------------------+ |',
    '|  |  COOLIFY-PROXY (Traefik) – Port 80/443                 | |',
    '|  |  SSL/TLS via Let\'s Encrypt (automatisch)               | |',
    '|  +-------------+----------------------------+-------------+ |',
    '|                |                            |               |',
    '|  +-------------+----------+  +--------------+-----------+   |',
    '|  |  FRONTEND  :3000        |  |  BACKEND  :8000          |  |',
    '|  |  Next.js 14 (Node 20)  |  |  FastAPI (Python 3.11)   |  |',
    '|  |  TailwindCSS           |  |  Anthropic SDK           |  |',
    '|  |  Leaflet.js            |  |  SSE-Streaming           |  |',
    '|  |  Static JSON Data      |  |  ZEO System-Prompt       |  |',
    '|  +------------------------+  +--------------------------+   |',
    '|                                                            |',
    '|  +--------------------------------------------------------+ |',
    '|  |  ANTHROPIC CLOUD API (extern)                          | |',
    '|  |  Claude claude-sonnet-4-6 – KI-Modell                         | |',
    '|  +--------------------------------------------------------+ |',
    '+============================================================+',
])

add_body(doc, '')
add_body(doc, 'Verwendete Technologien im Ueberblick:')

# Tabelle: Technologien
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
hdr[0].text = 'Komponente'
hdr[1].text = 'Technologie'
hdr[2].text = 'Begruendung'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ('Frontend',       'Next.js 14, App Router, TailwindCSS', 'Identisch mit mopilot.website – Wiederverwendung von Know-how'),
    ('Karte',          'Leaflet.js + OpenStreetMap',           'Keine API-Kosten, Open Source, datenschutzkonform'),
    ('Backend',        'FastAPI (Python 3.11)',                 'Leichtgewichtig, SSE-Streaming nativ unterstuetzt'),
    ('KI-Modell',      'Claude claude-sonnet-4-6',                     'Hohe Qualitaet, deutschsprachig, Anthropic API'),
    ('Datenhaltung',   'Statische JSON-Dateien',               'Kein Datenbank-Overhead fuer Prototyp'),
    ('Deployment',     'Docker Compose + Traefik-Labels',      'Integration in bestehende Coolify-Infrastruktur'),
    ('SSL',            'Let\'s Encrypt (via Traefik)',         'Automatisch, kostenlos, kein Wartungsaufwand'),
]
for r in rows:
    row = tbl.add_row().cells
    row[0].text, row[1].text, row[2].text = r

add_body(doc, '')

# 5.3
add_h2(doc, '5.3 Entwicklungsprozess mit Claude Code')
add_body(doc,
    'Die gesamte Applikation wurde in einer einzigen Session mit Claude Code entwickelt – '
    'dem KI-gestuetzten Entwicklungsassistenten in der lokalen Entwicklungsumgebung. '
    'Der Prozess war vollstaendig iterativ und dokumentengetrieben:')

add_h3(doc, 'Anforderungsanalyse')
add_body(doc,
    'Aus den vier PDF-Dateien im Ordner MoPilot_ZEO_Kunden\\ (Startbildschirm, Fahrzeugmodelle, '
    'Tarifuebersicht, FAQ) wurden alle relevanten ZEO-Inhalte extrahiert. Claude Code las die '
    'PDFs direkt mit pdftotext und strukturierte die Daten in JSON-Dateien:')
add_bullet(doc, 'vehicles.json: 9 Fahrzeugmodelle mit Specs (Sitze, Reichweite, Kofferraum, Ausstattung)')
add_bullet(doc, 'tariffs.json: Beide Tarife (eco, eco plus) mit allen Preisdetails')
add_bullet(doc, 'stations.json: 12 Musterstationen mit Koordinaten und ÖPNV-Anbindung')
add_bullet(doc, 'faq.json: 18 haeufige Fragen in 3 Kategorien')

add_h3(doc, 'Code-Generierung')
add_body(doc,
    'Ausgehend von den Anforderungen und der bestehenden MoPilot-Architektur generierte Claude Code '
    'alle Projektdateien in einem Schritt: Backend (FastAPI), Frontend (Next.js mit 8 Komponenten), '
    'Docker-Konfiguration und Nginx-Referenzkonfiguration. Insgesamt wurden 35 Dateien mit '
    '2.736 Codezeilen erstellt.')

add_h3(doc, 'Git-Workflow')
add_body(doc,
    'Der Code wurde auf einem neuen Feature-Branch (feature/zeo-kundenassistent) committet, '
    'ein Pull Request erstellt und nach Review in den master-Branch gemergt. '
    'Der Git-Commit-Prozess wurde vollstaendig von Claude Code ausgefuehrt.')

# 5.4
add_h2(doc, '5.4 Deployment-Ablauf: Herausforderungen und Loesungen')
add_body(doc,
    'Das Deployment auf die bestehende Hetzner-Infrastruktur verlief in mehreren Phasen '
    'und erforderte pragmatische Loesungen fuer auftretende Probleme:')

add_h3(doc, 'Phase 1: Coolify GitHub-Integration (Problem)')
add_body(doc,
    'Erster Versuch: Neuen Docker Compose Service in Coolify anlegen mit dem privaten '
    'GitHub-Repository als Quelle.')
add_body(doc, 'Problem:')
add_mono(doc, 'Failed to read Git source: fatal: could not read Username for')
add_mono(doc, "  'https://github.com': No such device or address")
add_body(doc,
    'Ursache: Das GitHub-Repository ist privat. Coolify benoetigt entweder eine '
    'GitHub App-Integration oder einen Personal Access Token (PAT). '
    'Die GitHub App-Konfiguration (App ID, Installation ID, Private Key) erwies sich '
    'als zu komplex fuer den Prototypen-Kontext.')

add_h3(doc, 'Phase 2: Direktes Deployment per SSH/SCP (Loesung)')
add_body(doc,
    'Da die Projektdateien lokal auf dem Entwickler-PC vorlagen (C:\\Projekte\\MoPilot\\MoPilot_ZEO_Kunden), '
    'wurde der Deployment-Weg gewechselt: direkter Upload per SFTP via Python-Paramiko-Bibliothek, '
    'ohne GitHub-Abhaengigkeit.')
add_body(doc, 'Vorgehen:')
add_mono(doc, '# Python-Script mit paramiko:')
add_mono(doc, '1. SSH-Verbindung zum Server herstellen')
add_mono(doc, '2. Alle 31 Projektdateien per SFTP auf /opt/zeo-kunden/ uebertragen')
add_mono(doc, '3. .env-Datei mit ANTHROPIC_API_KEY auf dem Server erstellen')
add_mono(doc, '4. docker compose up -d --build ausfuehren')

add_h3(doc, 'Phase 3: Traefik-Integration (Reverse Proxy)')
add_body(doc,
    'Die Containers wurden initial ohne Verbindung zum Coolify-Netzwerk gestartet, '
    'wodurch der Traefik-Proxy (coolify-proxy) sie nicht routen konnte. '
    'Loesung: docker-compose.yml um Traefik-Labels und externe Coolify-Netzwerkanbindung erweitert:')
add_mono(doc, 'networks:')
add_mono(doc, '  coolify:')
add_mono(doc, '    external: true   # Verbindung zum Traefik-Netzwerk')
add_mono(doc, '')
add_mono(doc, 'labels:')
add_mono(doc, '  - "traefik.enable=true"')
add_mono(doc, '  - "traefik.http.routers.zeo-frontend.rule=Host(`zeo-kunden.mopilot.website`)"')
add_mono(doc, '  - "traefik.http.routers.zeo-frontend.tls.certresolver=letsencrypt"')
add_mono(doc, '  - "traefik.http.services.zeo-frontend.loadbalancer.server.port=3000"')
add_body(doc,
    'Durch die Traefik-Labels erhielten die Container automatisch ein SSL-Zertifikat '
    'via Let\'s Encrypt und wurden ohne zusaetzliche Nginx-Konfiguration oeffentlich erreichbar.')

add_h3(doc, 'Phase 4: DNS-Konfiguration')
add_body(doc,
    'DNS fuer mopilot.website wird ueber IONOS verwaltet. Zwei neue A-Records wurden angelegt:')
add_mono(doc, 'zeo-kunden.mopilot.website     A  142.132.232.211  TTL: 1 Min.')
add_mono(doc, 'api.zeo-kunden.mopilot.website A  142.132.232.211  TTL: 1 Min.')
add_body(doc, 'Propagation: < 2 Minuten (durch kurze TTL von 1 Minute).')

add_h3(doc, 'Phase 5: API-Key-Rotation')
add_body(doc,
    'Nach dem ersten erfolgreichen Deployment wurde der im Deployment-Prozess verwendete '
    'Anthropic API-Key rotiert und der neue Key per SSH-Script in der .env-Datei auf dem '
    'Server aktualisiert. Nur das Backend wurde neu gestartet (kein Rebuild noetig):')
add_mono(doc, 'docker compose restart backend')

# 5.5
add_h2(doc, '5.5 Live-System und Ergebnis')
add_body(doc,
    'Das System ist seit dem 1. Maerz 2026 produktiv in Betrieb. '
    'Der Health-Check des Backends bestaetigt den erfolgreichen Start:')
add_mono(doc, '$ curl https://api.zeo-kunden.mopilot.website/api/health')
add_mono(doc, '{"status": "ok", "service": "zeo-kundenassistent"}')
add_body(doc, '')
add_body(doc, 'Beide Container laufen stabil im Coolify-Netzwerk:')
add_mono(doc, 'NAME                            IMAGE                    STATUS')
add_mono(doc, 'mopilot_zeo_kunden-frontend-1   zeo_kunden-frontend      Up')
add_mono(doc, 'mopilot_zeo_kunden-backend-1    zeo_kunden-backend       Up')
add_body(doc, '')
add_body(doc, 'Projektpfade im Ueberblick:')

tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = 'Table Grid'
h2 = tbl2.rows[0].cells
h2[0].text, h2[1].text = 'Ort', 'Pfad / URL'
for cell in h2:
    for run in cell.paragraphs[0].runs:
        run.bold = True
for r in [
    ('Lokal (Entwicklung)', 'C:\\Projekte\\MoPilot\\MoPilot_ZEO_Kunden\\'),
    ('GitHub (Versionierung)', 'github.com/JamesMz05/mopilot (Branch: master)'),
    ('Server (Deployment)', '/opt/zeo-kunden/MoPilot_ZEO_Kunden/'),
    ('Frontend (live)', 'https://zeo-kunden.mopilot.website'),
    ('Backend/API (live)', 'https://api.zeo-kunden.mopilot.website/api/health'),
]:
    row = tbl2.add_row().cells
    row[0].text, row[1].text = r

add_body(doc, '')

# ─── Kapitel 6: hotline.mopilot.website ────────────────────────────────────

add_h1(doc, '6. Naechstes Sub-Projekt: hotline.mopilot.website')
add_body(doc,
    'Aufbauend auf den Erfahrungen aus der ZEO-Kunden-App soll als naechstes Sub-Projekt '
    'ein spezialisierter Assistent fuer die ZEO Carsharing Hotline entwickelt werden. '
    'Dieses Kapitel beschreibt Konzept, Anforderungen, Entwicklungsschritte und Zeitplan.')

# 6.1
add_h2(doc, '6.1 Konzept und Zielgruppe')
add_body(doc,
    'hotline.mopilot.website richtet sich an die Mitarbeiterinnen und Mitarbeiter der '
    'ZEO Carsharing Hotline (betrieben durch Vianova eG). Im Gegensatz zum oeffentlichen '
    'ZEO Kundenassistenten ist dieses Tool ein internes Arbeitsmittel:')
add_bullet(doc, 'Zielgruppe: Hotline-Mitarbeiter (kein oeffentlicher Zugang, kein Login fuer Endkunden)')
add_bullet(doc, 'Kernfunktion: KI-gestuetzter Gespraeches-Assistent, der bei eingehenden Anrufen in Echtzeit unterstuetzt')
add_bullet(doc, 'Wissensquelle: Vollstaendige ZEO-Wissensbasis (Tarife, Fahrzeuge, Stationen, Prozesse, Sonderfaelle)')
add_bullet(doc, 'Zusatzfunktionen: Gespraechsleitfaden, Eskalationspfade, haeufige Problemszenarien')
add_body(doc,
    'Der Assistent fungiert als "zweite Stimme" im Ohr des Hotline-Mitarbeiters: '
    'waehrend der Mitarbeiter mit dem Kunden spricht, kann er parallel Fragen eingeben '
    'und sofortige, praezise Antworten erhalten – ohne lange in Handbuecher suchen zu muessen.')

# 6.2
add_h2(doc, '6.2 Technische Anforderungen')
add_body(doc, 'Die technischen Anforderungen unterscheiden sich in einigen Punkten von der ZEO-Kunden-App:')

tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = 'Table Grid'
h3 = tbl3.rows[0].cells
h3[0].text, h3[1].text, h3[2].text = 'Anforderung', 'Beschreibung', 'Prioritaet'
for cell in h3:
    for run in cell.paragraphs[0].runs:
        run.bold = True
for r in [
    ('Zugangsschutz',       'Basic Auth oder einfacher Login-Schutz (intern, kein oeffentlicher Zugang)', 'Hoch'),
    ('Schnelle Antworten',  'Antwortzeit < 3 Sekunden, SSE-Streaming fuer fliessende Ausgabe',            'Hoch'),
    ('Gespraeches-Kontext', 'Kontextgedaechtnis: laufendes Gespraech wird als Session gespeichert',       'Hoch'),
    ('Geraeteoptimiertheit','Optimiert fuer Desktop-Browser (Headset-Arbeitsplatz)',                      'Mittel'),
    ('Offline-Fallback',    'Wichtigste Infos auch ohne Internet abrufbar (statische Seite)',             'Niedrig'),
    ('Datenaktualitaet',    'Einfache Pflegbarkeit der Wissensbasis (JSON-Dateien im Git)',               'Mittel'),
]:
    row = tbl3.add_row().cells
    row[0].text, row[1].text, row[2].text = r

add_body(doc, '')
add_body(doc, 'Stack-Empfehlung: Identisch mit ZEO-Kunden-App (Next.js + FastAPI), '
    'ergaenzt um eine Session-State-Verwaltung im Frontend fuer den Gespraechsverlauf.')

# 6.3
add_h2(doc, '6.3 Entwicklungsschritte')
add_body(doc, 'Die Entwicklung gliedert sich in fuenf Phasen, die aufeinander aufbauen:')

add_h3(doc, 'Schritt 1: Inhaltsvorbereitung und System-Prompt (0,5 Tage)')
add_bullet(doc, 'Vollstaendige ZEO-Wissensbasis zusammenstellen: Alle Tarife, Sonderfaelle, Eskalationspfade, interne Prozesse')
add_bullet(doc, 'Hotline-spezifischen System-Prompt entwickeln: Tonalitaet (sachlich, schnell, praezise), Rollenbewusstsein')
add_bullet(doc, 'Gespraeches-Szenarien definieren: Typische Anrufgruende (Buchungsprobleme, Fahrzeugdefekte, Stornierungen)')
add_bullet(doc, 'JSON-Datensaetze aus ZEO-Kunden-App uebernehmen und um Hotline-Inhalte erweitern')

add_h3(doc, 'Schritt 2: Projektstruktur anlegen (0,5 Tage)')
add_body(doc, 'Neues Verzeichnis im bestehenden Repository anlegen:')
add_mono(doc, 'C:\\Projekte\\MoPilot\\MoPilot_Hotline\\')
add_mono(doc, '  frontend/    # Next.js 14')
add_mono(doc, '  backend/     # FastAPI (aus ZEO-App kopieren, System-Prompt anpassen)')
add_mono(doc, '  docker-compose.yml')
add_mono(doc, '  .env.example')
add_bullet(doc, 'Aus ZEO-Kunden-App: backend/main.py, Dockerfiles, docker-compose.yml kopieren')
add_bullet(doc, 'Neuer Subdomain-Eintrag: hotline.mopilot.website und api.hotline.mopilot.website')

add_h3(doc, 'Schritt 3: Frontend-Entwicklung (1-2 Tage)')
add_body(doc, 'Das Frontend unterscheidet sich von der ZEO-Kunden-App durch seinen Fokus auf Effizienz:')
add_bullet(doc, 'Kompaktes, zweispaltiges Layout: Links Chat, rechts Schnellreferenz-Panel')
add_bullet(doc, 'Schnellzugriff-Buttons: haeufigste Anrufgruende als vorgefertigte Fragen')
add_bullet(doc, 'Session-Protokoll: laufendes Gespraech wird angezeigt und kann exportiert werden')
add_bullet(doc, 'Gespraeches-Leitfaden: interaktiver Schritt-fuer-Schritt-Guide fuer Standard-Szenarien')
add_bullet(doc, 'Dunkel-Modus: fuer Bildschirmarbeit am Headset-Arbeitsplatz optimiert')
add_bullet(doc, 'Basic Auth Schutz via Next.js Middleware (einfache Passwortsicherung)')

add_h3(doc, 'Schritt 4: Backend-Anpassung (0,5 Tage)')
add_bullet(doc, 'System-Prompt auf Hotline-Rolle anpassen: Kuerze, Praezision, Eskalationshinweise')
add_bullet(doc, 'CORS-Origins auf hotline.mopilot.website beschraenken')
add_bullet(doc, 'Optional: Session-ID im Request mitfuehren fuer Gespraechsgedaechtnis')

add_h3(doc, 'Schritt 5: Deployment (0,5 Tage)')
add_body(doc, 'Deployment auf bestehende Infrastruktur – analog ZEO-Kunden-App:')
add_mono(doc, '1. DNS bei IONOS: hotline.mopilot.website  A  142.132.232.211')
add_mono(doc, '                  api.hotline.mopilot.website A 142.132.232.211')
add_mono(doc, '')
add_mono(doc, '2. Dateien per SFTP auf Server uebertragen:')
add_mono(doc, '   /opt/hotline/MoPilot_Hotline/')
add_mono(doc, '')
add_mono(doc, '3. docker-compose.yml: Traefik-Labels setzen (analog ZEO-Kunden):')
add_mono(doc, '   traefik.http.routers.hotline-frontend.rule=Host(`hotline.mopilot.website`)')
add_mono(doc, '')
add_mono(doc, '4. docker compose up -d --build')
add_mono(doc, '')
add_mono(doc, '5. SSL automatisch via Traefik/Let\'s Encrypt')

# 6.4
add_h2(doc, '6.4 Deployment-Checkliste')
add_body(doc, 'Vor dem Go-Live sind folgende Punkte abzuhaken:')

add_bullet(doc, 'Inhalt: ZEO-Wissensbasis vollstaendig und aktuell (Tarife, Prozesse, Eskalation)')
add_bullet(doc, 'Inhalt: System-Prompt mit Hotline-Mitarbeitern getestet und abgestimmt')
add_bullet(doc, 'Technik: DNS-Eintraege bei IONOS gesetzt und propagiert')
add_bullet(doc, 'Technik: Traefik-Labels in docker-compose.yml korrekt (Hostname, Port)')
add_bullet(doc, 'Technik: ANTHROPIC_API_KEY in .env gesetzt (separater Key empfohlen)')
add_bullet(doc, 'Technik: CORS_ORIGINS auf hotline.mopilot.website beschraenkt')
add_bullet(doc, 'Sicherheit: Basic Auth Passwort gesetzt und nur intern kommuniziert')
add_bullet(doc, 'Test: Chat-Funktion mit Beispiel-Anrufszenarien getestet')
add_bullet(doc, 'Test: SSL-Zertifikat aktiv (https:// erreichbar)')
add_bullet(doc, 'Test: API Health-Check positiv: /api/health -> {"status": "ok"}')
add_bullet(doc, 'Abnahme: Demo mit Hotline-Team durchgefuehrt und Feedback eingearbeitet')

# 6.5
add_h2(doc, '6.5 Zeitplan')
add_body(doc, 'Geschaetzter Gesamtaufwand: 3–4 Arbeitstage fuer Entwicklung und Deployment.')

tbl4 = doc.add_table(rows=1, cols=3)
tbl4.style = 'Table Grid'
h4 = tbl4.rows[0].cells
h4[0].text, h4[1].text, h4[2].text = 'Phase', 'Aufwand', 'Voraussetzung'
for cell in h4:
    for run in cell.paragraphs[0].runs:
        run.bold = True
for r in [
    ('Inhaltsvorbereitung & System-Prompt', '0,5 Tage', 'Abstimmung mit ZEO Hotline-Team'),
    ('Projektstruktur & Code-Basis',        '0,5 Tage', 'ZEO-Kunden-App als Vorlage'),
    ('Frontend-Entwicklung',                '1-2 Tage', 'Hotline-UX-Anforderungen bekannt'),
    ('Backend-Anpassung',                   '0,5 Tage', 'System-Prompt fertiggestellt'),
    ('Deployment & Test',                   '0,5 Tage', 'DNS & Server-Zugang verfuegbar'),
    ('GESAMT',                              '3-4 Tage', ''),
]:
    row = tbl4.add_row().cells
    row[0].text, row[1].text, row[2].text = r
    if r[0] == 'GESAMT':
        for cell in row:
            for run in cell.paragraphs[0].runs:
                run.bold = True

add_body(doc, '')
add_body(doc,
    'Durch die Wiederverwendung der ZEO-Kunden-App-Architektur (Docker Compose, Traefik-Labels, '
    'FastAPI-Backend, Next.js-Frontend) reduziert sich der Aufwand erheblich. '
    'Der kritische Pfad liegt in der Inhaltsvorbereitung und der Abstimmung des System-Prompts '
    'mit dem Hotline-Team – nicht in der technischen Umsetzung.')

# Footer-Zeile aktualisieren
for para in doc.paragraphs:
    if 'Erstellt mit Claude Code am 28. Februar 2026' in para.text:
        for run in para.runs:
            if '28. Februar 2026' in run.text:
                run.text = run.text.replace('28. Februar 2026', '1. Maerz 2026')
        # Version ersetzen
        if 'v0.1' in para.text:
            for run in para.runs:
                run.text = run.text.replace('v0.1', 'v0.2')

doc.save(DEST)
print(f"Gespeichert: {DEST}")
print(f"Kapitel 5 und 6 eingefuegt.")
